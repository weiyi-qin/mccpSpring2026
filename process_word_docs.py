#!/usr/bin/env python3
"""
Process Word documents from Materials/word with real-time logging and careful image handling.
Uses LLM (Poe API) to format markdown output.
Tracks processing status in CSV index file.
"""

import os
import sys
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from pathlib import Path
import requests
import json
import re
from datetime import datetime
import time
import csv
from io import BytesIO
from PIL import Image

# Global log file
LOG_FILE = None

def log(message, level="INFO"):
    """Write to both console and log file."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_message = f"[{timestamp}] [{level}] {message}"
    print(log_message)
    if LOG_FILE:
        LOG_FILE.write(log_message + "\n")
        LOG_FILE.flush()  # Ensure real-time writing

def get_poe_api_key():
    """Read Poe API key from LLM/poe.md"""
    poe_file = Path(__file__).parent / "LLM" / "poe.md"
    if poe_file.exists():
        with open(poe_file, 'r') as f:
            lines = [line.strip() for line in f.readlines() if line.strip() and not line.strip().startswith('http')]
            if lines:
                return lines[0]
    return None

def extract_images_from_word(docx_path, output_dir):
    """Extract images and charts from Word document."""
    images_info = []
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    log(f"Extracting images from: {docx_path.name}")
    
    try:
        doc = Document(docx_path)
        image_count = 0
        
        # Access the document's related parts (images)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Determine image format
                    ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                    if ext not in ['png', 'jpg', 'jpeg', 'gif']:
                        ext = 'png'
                    
                    # Save image
                    img_filename = f"image_{image_count:03d}.{ext}"
                    img_path = output_dir / img_filename
                    
                    with open(img_path, 'wb') as f:
                        f.write(image_data)
                    
                    # Get image dimensions if possible
                    try:
                        img = Image.open(BytesIO(image_data))
                        width, height = img.size
                        log(f"  Saved image {image_count}: {img_filename} ({width}x{height})")
                    except:
                        log(f"  Saved image {image_count}: {img_filename}")
                    
                    images_info.append({
                        'index': image_count,
                        'path': str(img_path),
                        'filename': img_filename,
                        'type': 'embedded'
                    })
                except Exception as e:
                    log(f"  Warning: Could not extract image {image_count}: {e}", "WARNING")
        
        # Also check inline shapes
        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:
                if run._element.xml.find('pic:pic') != -1:
                    # Found an inline image
                    pass  # Already handled via rels
        
        if not images_info:
            log(f"  No images found in document")
        else:
            log(f"  ✓ Extracted {len(images_info)} images")
        
        return images_info
    except Exception as e:
        log(f"Error extracting images: {e}", "ERROR")
        import traceback
        log(traceback.format_exc(), "ERROR")
        return []

def extract_text_from_word(docx_path):
    """Extract text from Word document with structure preservation."""
    text_content = []
    log(f"Extracting text from: {docx_path.name}")
    
    try:
        doc = Document(docx_path)
        full_text = []
        char_count = 0
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect heading level
                style_name = para.style.name if para.style else ""
                if style_name.startswith('Heading'):
                    level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                    full_text.append(f"\n{'#' * (level + 1)} {text}\n")
                elif para.style.name == 'List Bullet' or para.style.name == 'List Paragraph':
                    full_text.append(f"- {text}\n")
                elif para.style.name == 'List Number':
                    full_text.append(f"1. {text}\n")
                else:
                    full_text.append(f"{text}\n")
                char_count += len(text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            full_text.append(f"\n### Table {table_idx + 1}\n\n")
            for row in table.rows:
                row_text = "| " + " | ".join([cell.text.strip() for cell in row.cells]) + " |"
                full_text.append(row_text + "\n")
                if table_idx == 0 and row == table.rows[0]:  # Header row
                    full_text.append("|" + "|".join(["---"] * len(row.cells)) + "|\n")
            full_text.append("\n")
        
        combined_text = "".join(full_text)
        text_content = {
            'text': combined_text,
            'char_count': char_count,
            'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
            'table_count': len(doc.tables)
        }
        
        log(f"  Extracted {char_count} characters from {text_content['paragraph_count']} paragraphs and {text_content['table_count']} tables")
        
        return text_content
    except Exception as e:
        log(f"Error extracting text: {e}", "ERROR")
        import traceback
        log(traceback.format_exc(), "ERROR")
        return None

def call_poe_api(prompt, api_key, max_retries=2):
    """Call Poe API to process content with LLM."""
    if not api_key:
        log("No Poe API key found. Skipping LLM processing.", "WARNING")
        return None
    
    try:
        url = "https://api.poe.com/v1/chat/completions"
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "Claude-3.5-Sonnet",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": 16000,
            "temperature": 0.3
        }
        
        log(f"  Calling Poe API (model: Claude-3.5-Sonnet)...")
        start_time = time.time()
        
        for attempt in range(max_retries):
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=180)
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
                    elapsed = time.time() - start_time
                    log(f"  ✓ API call successful ({elapsed:.1f}s, {len(content)} chars returned)")
                    return content
                else:
                    log(f"  API error (attempt {attempt+1}/{max_retries}): {response.status_code} - {response.text[:200]}", "WARNING")
                    if attempt < max_retries - 1:
                        time.sleep(2)
            except requests.Timeout:
                log(f"  API timeout (attempt {attempt+1}/{max_retries})", "WARNING")
                if attempt < max_retries - 1:
                    time.sleep(2)
        
        return None
    except Exception as e:
        log(f"Error calling Poe API: {e}", "ERROR")
        return None

def format_markdown_with_llm(text_content, images_info, api_key, doc_name):
    """Use LLM to format markdown content properly."""
    log(f"Formatting markdown with LLM...")
    
    # Prepare content for LLM
    text = text_content['text']
    
    # Add image references
    image_refs = ""
    if images_info:
        image_refs = "\n\n## Images and Charts\n\n"
        for img_info in images_info:
            session_folder = Path(img_info['path']).parent.name
            rel_path = f"images/{session_folder}/{img_info['filename']}"
            image_refs += f"![Image {img_info['index']}]({rel_path})\n\n"
    
    # Create prompt
    prompt = f"""Format this Word document content into well-structured markdown. CRITICAL: Preserve ALL content exactly - do NOT summarize, abbreviate, or omit anything.

Requirements:
1. Use proper markdown headings (##, ###, ####)
2. Format lists properly (- for bullets, 1. for numbered)
3. Make task titles bold: **Task X**, **Warm-up Task**
4. Format URLs as [text](url)
5. Format tables properly (markdown table syntax)
6. Add proper spacing (--- for major breaks)
7. CRITICAL: Include ALL text - every word must be preserved
8. Do NOT add summaries, "[Continues...]", or placeholders
9. Preserve all examples, excerpts, citations exactly

Word Document Content:
{text[:15000]}

Return the complete formatted markdown with ALL content preserved:"""

    # Call LLM
    formatted = call_poe_api(prompt, api_key)
    
    if formatted:
        # Add image references at the end
        if image_refs:
            formatted += image_refs
        log(f"  ✓ LLM formatting complete ({len(formatted)} characters)")
        return formatted
    else:
        # Fallback to basic formatting
        log("  Using fallback formatting (no LLM)", "WARNING")
        return format_markdown_basic(text_content, images_info)

def format_markdown_basic(text_content, images_info):
    """Basic markdown formatting without LLM."""
    md_content = []
    md_content.append(text_content['text'])
    
    # Add images
    if images_info:
        md_content.append("\n\n---\n\n## Images and Charts\n\n")
        for img_info in images_info:
            session_folder = Path(img_info['path']).parent.name
            rel_path = f"images/{session_folder}/{img_info['filename']}"
            md_content.append(f"![Image {img_info['index']}]({rel_path})\n\n")
    
    return "\n".join(md_content)

def update_csv_index(csv_path, filename, status, text_extracted=None, images_extracted=None, markdown_file=None, notes=None):
    """Update the CSV index file with processing status."""
    try:
        # Read existing data
        rows = []
        headers = ['filename', 'file_type', 'status', 'processed_date', 'text_extracted', 'images_extracted', 'markdown_file', 'notes']
        
        if csv_path.exists():
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
        
        # Update or add row
        updated = False
        for row in rows:
            if row['filename'] == filename:
                row['status'] = status
                if text_extracted is not None:
                    row['text_extracted'] = str(text_extracted)
                if images_extracted is not None:
                    row['images_extracted'] = str(images_extracted)
                if markdown_file:
                    row['markdown_file'] = markdown_file
                if notes:
                    row['notes'] = notes
                if status == 'completed':
                    row['processed_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                updated = True
                break
        
        if not updated:
            # Add new row
            new_row = {
                'filename': filename,
                'file_type': 'docx',
                'status': status,
                'processed_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S") if status == 'completed' else '',
                'text_extracted': str(text_extracted) if text_extracted else '',
                'images_extracted': str(images_extracted) if images_extracted else '',
                'markdown_file': markdown_file or '',
                'notes': notes or ''
            }
            rows.append(new_row)
        
        # Write back
        with open(csv_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=headers)
            writer.writeheader()
            writer.writerows(rows)
        
        log(f"  Updated CSV index: {filename} -> {status}")
    except Exception as e:
        log(f"Error updating CSV index: {e}", "WARNING")

def process_single_word_doc(docx_path, output_md_dir, output_img_dir, api_key, csv_path):
    """Process a single Word document."""
    docx_path = Path(docx_path)
    log(f"\n{'='*70}")
    log(f"PROCESSING: {docx_path.name}")
    log(f"{'='*70}")
    
    # Update CSV: in_progress
    update_csv_index(csv_path, docx_path.name, 'in_progress')
    
    # Extract images
    session_name = docx_path.stem.replace('_TEACHER version', '').replace('_TEACHER Version', '') + '_word'
    img_output_dir = Path(output_img_dir) / session_name
    images_info = extract_images_from_word(docx_path, img_output_dir)
    log(f"✓ Extracted {len(images_info)} images")
    
    # Extract text
    text_content = extract_text_from_word(docx_path)
    if not text_content:
        log(f"✗ Failed to extract text from {docx_path.name}", "ERROR")
        update_csv_index(csv_path, docx_path.name, 'error', notes='Failed to extract text')
        return False
    
    log(f"✓ Extracted {text_content['char_count']} characters")
    update_csv_index(csv_path, docx_path.name, 'in_progress', 
                     text_extracted=text_content['char_count'], 
                     images_extracted=len(images_info))
    
    # Format with LLM
    md_content = format_markdown_with_llm(text_content, images_info, api_key, docx_path.stem)
    
    # Save markdown
    md_filename = docx_path.stem + "_word.md"
    md_path = Path(output_md_dir) / md_filename
    
    log(f"Saving markdown to: {md_path.name}")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    log(f"✓ Created markdown file: {md_path} ({len(md_content)} characters)")
    log(f"✓ Processed {len(images_info)} images to {img_output_dir}")
    
    # Update CSV: completed
    update_csv_index(csv_path, docx_path.name, 'completed',
                     text_extracted=text_content['char_count'],
                     images_extracted=len(images_info),
                     markdown_file=md_filename)
    
    return True

def get_user_feedback(doc_name, auto_continue=False):
    """Ask user for feedback before proceeding."""
    print(f"\n{'='*70}")
    print(f"Completed processing: {doc_name}")
    print(f"{'='*70}")
    
    if auto_continue:
        log("Auto-continuing to next document (non-interactive mode)")
        return '1'
    
    print("\nOptions:")
    print("1. Continue to next document")
    print("2. Review and provide feedback")
    print("3. Skip remaining documents")
    print("4. Exit")
    
    try:
        while True:
            choice = input("\nEnter your choice (1-4): ").strip()
            if choice in ['1', '2', '3', '4']:
                return choice
            print("Invalid choice. Please enter 1, 2, 3, or 4.")
    except (EOFError, KeyboardInterrupt):
        log("Non-interactive mode detected, auto-continuing", "WARNING")
        return '1'

def main():
    global LOG_FILE
    
    base_dir = Path("/Users/simonwang/Library/CloudStorage/OneDrive-HongKongBaptistUniversity/GTD/Areas/Teaching/Courses/MCCP 6020/PhDagentSpring2026")
    word_dir = base_dir / "Materials" / "word"
    md_dir = base_dir / "Materials" / "md"
    img_dir = base_dir / "Materials" / "md" / "images"
    csv_path = word_dir / "processing_index.csv"
    log_file_path = base_dir / "process_word.log"
    
    # Initialize log file
    LOG_FILE = open(log_file_path, 'w', encoding='utf-8')
    log(f"Starting Word document processing session")
    log(f"Word directory: {word_dir}")
    log(f"Output MD directory: {md_dir}")
    log(f"Output images directory: {img_dir}")
    log(f"CSV index: {csv_path}")
    
    # Ensure directories exist
    md_dir.mkdir(parents=True, exist_ok=True)
    img_dir.mkdir(parents=True, exist_ok=True)
    
    # Get API key
    api_key = get_poe_api_key()
    if api_key:
        log(f"✓ Found Poe API key")
    else:
        log("⚠ No Poe API key found - will use basic formatting", "WARNING")
    
    # Get all Word documents (skip PDFs for now)
    word_files = sorted([f for f in word_dir.glob("*.docx")])
    
    if not word_files:
        log(f"✗ No Word documents found in {word_dir}", "ERROR")
        LOG_FILE.close()
        return
    
    log(f"\nFound {len(word_files)} Word documents to process:")
    for i, word_file in enumerate(word_files, 1):
        log(f"  {i}. {word_file.name}")
    
    # Process each Word document one at a time
    processed = 0
    for i, word_file in enumerate(word_files, 1):
        log(f"\n{'#'*70}")
        log(f"Processing Word document {i}/{len(word_files)}")
        log(f"{'#'*70}")
        
        success = process_single_word_doc(word_file, md_dir, img_dir, api_key, csv_path)
        
        if success:
            processed += 1
            log(f"✓ Successfully processed: {word_file.name}")
        else:
            log(f"✗ Failed to process: {word_file.name}", "ERROR")
        
        # Ask for feedback
        if i < len(word_files):
            auto_mode = not sys.stdin.isatty()
            choice = get_user_feedback(word_file.name, auto_continue=auto_mode)
            
            if choice == '2':
                print("\nPlease review the output files and provide any feedback.")
                input("Press Enter when ready to continue...")
            elif choice == '3':
                log("User chose to skip remaining documents")
                break
            elif choice == '4':
                log("User chose to exit")
                break
    
    log(f"\n{'='*70}")
    log(f"Processing session complete")
    log(f"Processed: {processed}/{len(word_files)} Word documents")
    log(f"{'='*70}")
    
    LOG_FILE.close()
    print(f"\n✓ Processing complete! Check process_word.log for details.")
    print(f"✓ Status tracked in: {csv_path}")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        if LOG_FILE:
            log("\nProcess interrupted by user", "WARNING")
            LOG_FILE.close()
        sys.exit(1)
    except Exception as e:
        if LOG_FILE:
            log(f"Fatal error: {e}", "ERROR")
            import traceback
            log(traceback.format_exc(), "ERROR")
            LOG_FILE.close()
        raise
