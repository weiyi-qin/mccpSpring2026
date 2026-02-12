#!/usr/bin/env python3
"""
Improve markdown formatting by directly extracting from Word with proper structure,
then using LLM only for final polish (not content generation).
"""

import os
import sys
from docx import Document
from pathlib import Path
import requests
import re
from datetime import datetime

def get_poe_api_key():
    """Read Poe API key from LLM/poe.md"""
    poe_file = Path(__file__).parent.parent / "LLM" / "poe.md"
    if poe_file.exists():
        with open(poe_file, 'r') as f:
            lines = [line.strip() for line in f.readlines() if line.strip() and not line.strip().startswith('http')]
            if lines:
                return lines[0]
    return None

def format_table_markdown(table_data):
    """Convert table data to markdown table format."""
    if not table_data or not table_data[0]:
        return ""
    
    md_lines = []
    
    # Header row
    header = table_data[0]
    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("|" + "|".join(["---"] * len(header)) + "|")
    
    # Data rows
    for row in table_data[1:]:
        if row and any(cell.strip() for cell in row):
            # Pad row if needed
            while len(row) < len(header):
                row.append("")
            md_lines.append("| " + " | ".join(row[:len(header)]) + " |")
    
    return "\n".join(md_lines)

def extract_and_format_word_document(docx_path):
    """Extract and format Word document directly to markdown."""
    doc = Document(docx_path)
    md_parts = []
    
    print(f"Extracting from: {Path(docx_path).name}")
    print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    
    # Extract session number
    session_match = re.search(r'Session (\d+)', Path(docx_path).stem)
    session_num = session_match.group(1) if session_match else "X"
    
    # Add title
    md_parts.append(f"# MCCP6020 ADVANCED ENGLISH FOR ACADEMIC PURPOSES SESSION {session_num}\n")
    md_parts.append("---\n\n")
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else "Normal"
        
        # Detect and format structure
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
            except:
                level = 1
            md_parts.append(f"\n{'#' * (level + 1)} {text}\n")
        elif style_name in ['List Paragraph', 'List Bullet']:
            md_parts.append(f"- {text}\n")
        elif style_name == 'List Number':
            md_parts.append(f"1. {text}\n")
        elif re.match(r'^(Task \d+|Warm-up Task)', text, re.IGNORECASE):
            # Make task titles bold
            md_parts.append(f"\n**{text}**\n")
        elif text == 'Learning Outcomes:':
            md_parts.append(f"\n**{text}**\n")
        elif 'Note to Teachers' in text:
            md_parts.append(f"\n> **Note to Teachers:** {text.replace('Note to Teachers:', '').strip()}\n")
        elif text == 'Introducing Your Research':
            # Section heading
            md_parts.append(f"\n---\n\n## {text}\n")
        elif text in ['Structure of a Thesis/Journal Article', 'The \'Hourglass\' Thesis Model Based on Primary Research', 
                      'Claiming Centrality of Your Research Topic', 'Doing Things with Nouns (Nominalization)', 
                      'Independent Language Learning', 'INTRODUCTION of a Thesis/Journal Article',
                      'The Creating a Research Space (CARS) model']:
            # Major section headings - exact matches only
            md_parts.append(f"\n---\n\n## {text}\n")
        elif text.startswith('The \'Hourglass\''):
            # Subsection
            md_parts.append(f"\n### {text}\n")
        else:
            # Regular paragraph
            md_parts.append(f"{text}\n")
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        if table_data:
            md_parts.append(f"\n### Table {table_idx + 1}\n\n")
            md_parts.append(format_table_markdown(table_data))
            md_parts.append("\n\n")
    
    # Add images section
    images_dir = Path(docx_path).parent.parent / "md" / "images" / f"Session {session_num}_Handout_word"
    if images_dir.exists():
        image_files = sorted(list(images_dir.glob("*.png")) + list(images_dir.glob("*.jpeg")) + list(images_dir.glob("*.jpg")))
        if image_files:
            md_parts.append("\n---\n\n## Images and Charts\n\n")
            for img_file in image_files:
                rel_path = f"images/Session {session_num}_Handout_word/{img_file.name}"
                md_parts.append(f"![{img_file.stem}]({rel_path})\n\n")
    
    return "".join(md_parts)

def polish_with_llm(markdown_content, api_key, session_num):
    """Use LLM to polish formatting without changing content."""
    if not api_key:
        return markdown_content
    
    # Process in chunks if too long
    if len(markdown_content) > 15000:
        print("  Content large, polishing in sections...")
        # Split by major sections (## headings)
        sections = re.split(r'\n(## .+)\n', markdown_content)
        polished_parts = []
        
        for i in range(0, len(sections), 2):
            if i < len(sections):
                section_header = sections[i] if i == 0 else f"\n{sections[i-1]}\n"
                section_content = sections[i] if i == 0 else sections[i]
                
                if len(section_content) > 8000:
                    # Section too large, use as-is
                    polished_parts.append(section_header + section_content)
                else:
                    # Polish this section
                    prompt = f"""Improve markdown formatting. CRITICAL: Do NOT change or remove content, only improve formatting.

Fix:
- Ensure proper heading levels
- Make task titles bold: **Task X**
- Format URLs: [text](url)
- Add spacing between sections
- Fix table formatting

Content:
{section_content[:8000]}

Return improved formatting with ALL content preserved:"""
                    
                    polished_section = call_poe_api_simple(prompt, api_key)
                    if polished_section:
                        polished_parts.append(section_header + polished_section)
                    else:
                        polished_parts.append(section_header + section_content)
        
        return "".join(polished_parts)
    
    # Process entire document if small enough
    
    prompt = f"""Improve the markdown formatting of this document. CRITICAL: Do NOT change, remove, or summarize any content. Only improve formatting.

What to do:
- Ensure consistent heading levels (## for main sections, ### for subsections)
- Fix any markdown syntax issues
- Ensure proper spacing between sections
- Make sure task titles are bold: **Task X**, **Warm-up Task**
- Make "Learning Outcomes:" bold: **Learning Outcomes:**
- Ensure URLs are formatted as [text](url)
- Fix any table formatting issues
- Add horizontal rules (---) between major sections if missing
- Format section titles like "Structure of a Thesis/Journal Article" as ## headings
- Format "Introducing Your Research" as ## heading

What NOT to do:
- Do NOT remove any text
- Do NOT summarize content
- Do NOT add "[Continues...]" or placeholders
- Do NOT change the actual content

Current markdown:
{markdown_content[:15000]}

Return the improved markdown with better formatting but ALL content preserved:"""

    try:
        url = "https://api.poe.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "Claude-3.5-Sonnet",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 20000,
            "temperature": 0.2
        }
        
        print("  Polishing with LLM...")
        response = requests.post(url, headers=headers, json=payload, timeout=180)
        
        if response.status_code == 200:
            result = response.json()
            polished = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            # Remove any LLM prefixes
            polished = re.sub(r'^[^\#]*', '', polished, flags=re.MULTILINE)
            print(f"  ✓ LLM polish complete ({len(polished)} chars)")
            return polished.strip()
        else:
            print(f"  LLM polish failed, using direct extraction")
            return markdown_content
    except Exception as e:
        print(f"  LLM polish error: {e}, using direct extraction")
        return markdown_content

def call_poe_api_simple(prompt, api_key):
    """Simple Poe API call."""
    try:
        url = "https://api.poe.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "Claude-3.5-Sonnet",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 10000,
            "temperature": 0.2
        }
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('message', {}).get('content', '')
        return None
    except:
        return None

def main():
    base_dir = Path("/Users/simonwang/Library/CloudStorage/OneDrive-HongKongBaptistUniversity/GTD/Areas/Teaching/Courses/MCCP 6020/PhDagentSpring2026")
    word_dir = base_dir / "Materials" / "word"
    md_dir = base_dir / "Materials" / "md"
    
    # Get API key
    api_key = get_poe_api_key()
    
    # Find all Word documents
    word_files = sorted(list(word_dir.glob("*.docx")))
    
    if not word_files:
        print(f"✗ No Word files found in: {word_dir}")
        return
    
    print(f"\n{'='*70}")
    print(f"Improving markdown formatting for {len(word_files)} Word documents")
    print(f"{'='*70}\n")
    
    for docx_file in word_files:
        # Generate output filename
        md_filename = docx_file.stem + "_word.md"
        md_file = md_dir / md_filename
        
        print(f"\n{'='*70}")
        print(f"Processing: {docx_file.name}")
        print(f"{'='*70}")
        
        # Extract directly from Word
        markdown_content = extract_and_format_word_document(docx_file)
        print(f"  ✓ Direct extraction: {len(markdown_content)} characters")
        
        # Save
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"✓ Saved improved markdown: {md_file.name}")
        print(f"  Final length: {len(markdown_content)} characters")
    
    print(f"\n{'='*70}")
    print(f"✓ Formatting improvement complete for all {len(word_files)} documents!")
    print(f"{'='*70}")

if __name__ == "__main__":
    main()
